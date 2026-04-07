"""Generate the Red Tail Surveying Workflow Reference as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title Page ──
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Red Tail Surveying')
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x2D, 0x8A, 0x6E)
r.bold = True

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Job Workflow Reference')
r2.font.size = Pt(22)
r2.font.color.rgb = RGBColor(0x4F, 0xAC, 0xFE)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Autonomous Research Pipeline — Team Review Document')
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()
legend = doc.add_paragraph()
legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = legend.add_run('GREEN = AI handles automatically     YELLOW = AI assists, human reviews     RED = Human required')
r.font.size = Pt(10)

doc.add_page_break()

# ── Quick Reference Table ──
doc.add_heading('Quick Reference: Job Types', level=1)
doc.add_paragraph('Overview of all job types and their research requirements.')

cols = ['Code', 'Type', 'Adjoiners?', 'Field Work', 'Plat?', 'Duration']
data = [
    ['BDY-E', 'Boundary - Existing Corners', 'Yes (full)', 'Monument search', 'Yes - retracement', '2-4 days'],
    ['BDY-N', 'Boundary - New Corners', 'Yes (full)', 'Set monuments', 'Yes - original', '3-6 days'],
    ['FT', 'FEMA Topo', 'No', 'Topo shoot', 'FEMA cert only', '1-2 days'],
    ['SUB', 'Subdivision', 'Yes (perimeter)', 'Set lot corners', 'Yes - subdivision', '5-10 days'],
    ['LLA', 'Lot Line Adjustment', 'Yes (affected)', 'Set new line', 'Yes - LLA plat', '2-4 days'],
    ['LC', 'Land Consolidation', 'Yes (perimeter)', 'Verify corners', 'Yes - consolidation', '2-3 days'],
    ['EAS', 'Easement', 'Partial', 'Locate/set', 'Yes - exhibit', '1-3 days'],
    ['P', 'Permit Survey', 'No', 'Site verification', 'Site plan only', '0.5-1 day'],
    ['SE', 'Survey Exam', 'Yes (full)', 'Review only', 'Report only', '1-2 days'],
    ['ILR', 'Improvement Location', 'No', 'Measure structures', 'ILR cert', '0.5-1 day'],
    ['WR', 'Water Rights', 'No', 'Ditch mapping', 'Exhibit map', '1-2 days'],
]

table = doc.add_table(rows=1, cols=len(cols))
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, c in enumerate(cols):
    table.rows[0].cells[i].text = c
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_page_break()

# ── Workflow helper ──
def add_workflow(title, scenario, research_steps, field_steps, drafting_steps, tip=None, important=None):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    r = p.add_run('Scenario: ')
    r.bold = True
    p.add_run(scenario)
    doc.add_paragraph()

    # Research
    doc.add_heading('Research Phase (Office)', level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Step', 'Task', 'Who', 'Notes']):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for s in research_steps:
        row = t.add_row()
        for i, v in enumerate(s):
            row.cells[i].text = str(v)

    # Field
    if field_steps:
        doc.add_heading('Field Phase', level=2)
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = 'Light Shading Accent 1'
        for i, h in enumerate(['Step', 'Task', 'Who', 'Notes']):
            t2.rows[0].cells[i].text = h
            t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for s in field_steps:
            row = t2.add_row()
            for i, v in enumerate(s):
                row.cells[i].text = str(v)

    # Drafting
    if drafting_steps:
        doc.add_heading('Drafting Phase', level=2)
        t3 = doc.add_table(rows=1, cols=4)
        t3.style = 'Light Shading Accent 1'
        for i, h in enumerate(['Step', 'Task', 'Who', 'Notes']):
            t3.rows[0].cells[i].text = h
            t3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for s in drafting_steps:
            row = t3.add_row()
            for i, v in enumerate(s):
                row.cells[i].text = str(v)

    if tip:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run('TIP - Team Customization: ')
        r.bold = True
        r.font.color.rgb = RGBColor(0x2D, 0x8A, 0x6E)
        p.add_run(tip)

    if important:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run('IMPORTANT: ')
        r.bold = True
        r.font.color.rgb = RGBColor(0xDA, 0x36, 0x33)
        p.add_run(important)

    doc.add_page_break()


# ── BDY-E ──
add_workflow(
    'Workflow 1: BDY-E - Boundary Survey (Existing Corners)',
    'Client owns property, corners have been set by a prior survey. Retracing and verifying the existing boundary.',
    [
        ['1', 'AI - Create project folders & session', 'AI', 'Auto folder tree'],
        ['2', 'AI - Search client deed (grantor index)', 'AI', 'Auto-select if confidence >= 25'],
        ['3', 'REVIEW - Verify correct parcel', 'Human', 'AI shows top matches, you confirm'],
        ['4', 'AI - Extract legal description', 'AI', 'Metes & bounds, lot/block, or TRS'],
        ['5', 'AI - Search client plat (cabinet)', 'AI', 'Matches by name in cabinet files'],
        ['6', 'REVIEW - Check plat for prior survey', 'Human', 'Look for existing monuments noted'],
        ['7', 'AI - Discover adjoiners (ArcGIS)', 'AI', '500m radius, touching parcels'],
        ['8', 'AI - Search adjoiner deeds (bulk)', 'AI', 'Auto-select best deed per adjoiner'],
        ['9', 'AI - Search adjoiner plats (cabinet)', 'AI', 'Cabinet match by last name'],
        ['10', 'REVIEW - Check adjoiner deeds', 'Human', 'Check calls for consistency'],
        ['11', 'AI - Build reference table', 'AI', 'All deeds & plats cited on plat'],
        ['12', 'AI - Generate DXF boundary', 'AI', 'If metes & bounds available'],
    ],
    [
        ['13', 'FIELD - Locate existing monuments', 'Surveyor', 'Rebar, pipe, stone - per plat'],
        ['14', 'FIELD - Survey control / traverse', 'Surveyor', 'Control network for boundary'],
        ['15', 'FIELD - Measure to all corners', 'Surveyor', 'Compare to record distances'],
        ['16', 'FIELD - Photograph monuments', 'Surveyor', 'Evidence file'],
        ['17', 'FIELD - Resolve discrepancies', 'PLS', 'Professional opinion if corners conflict'],
    ],
    [
        ['18', 'AI+HUMAN - Generate boundary plat', 'AI+Human', 'AI creates DXF, human refines'],
        ['19', 'AI - Insert reference table', 'AI', 'Auto-generated from research'],
        ['20', 'HUMAN - Surveyors certificate', 'PLS', 'Legal certification'],
        ['21', 'HUMAN - Sign & seal', 'PLS', 'Licensed surveyor signature'],
        ['22', 'AI - File to project folder', 'AI', 'Auto-organize output files'],
    ],
    tip='Add your specific monument search protocol here. Do you always check PLSS corners first? Do you have a preferred order for checking adjoiners corners?',
)

# ── BDY-N ──
add_workflow(
    'Workflow 2: BDY-N - Boundary Survey (New Corners)',
    'Property has never been surveyed, or prior survey corners are lost/destroyed. Surveyor must establish new corners.',
    [
        ['1-12', 'Same as BDY-E above', 'AI/Human', 'Full research pipeline'],
        ['12a', 'REVIEW - Analyze legal desc for ambiguity', 'AI+Human', 'AI flags missing calls, unclosed traverses'],
        ['12b', 'HUMAN - Determine corner method', 'PLS', 'Proportion, lost/obliterated rules'],
        ['12c', 'AI - Check chain of title', 'AI', 'Chain search + prior deed analysis'],
        ['12d', 'HUMAN - Search for BLM/GLO plats', 'Human', 'Federal survey records if PLSS'],
    ],
    [
        ['13', 'FIELD - Search for existing evidence', 'Surveyor', 'Fences, walls, occupation lines'],
        ['14', 'FIELD - Survey control / traverse', 'Surveyor', 'Higher precision - setting corners'],
        ['15', 'FIELD - Calculate new corner positions', 'PLS', 'Based on deed, adjoiners, evidence'],
        ['16', 'FIELD - Set rebar & caps at all corners', 'Surveyor', 'Monument with PLS number'],
        ['17', 'FIELD - Tie to PLSS or reference', 'Surveyor', 'For future recovery'],
        ['18', 'FIELD - Photograph all set monuments', 'Surveyor', 'Evidence file'],
    ],
    [
        ['19', 'AI+HUMAN - Generate boundary plat', 'AI+Human', 'Must show NEW corners distinctly'],
        ['20', 'HUMAN - Add monument descriptions', 'PLS', 'Set 1/2" rebar with RTS cap'],
        ['21', 'HUMAN - Add basis of bearing', 'PLS', 'How corners were established'],
        ['22', 'AI - Insert reference table', 'AI', ''],
        ['23', 'HUMAN - Certificate (original survey)', 'PLS', 'Different language than retracement'],
        ['24', 'HUMAN - Sign & seal', 'PLS', ''],
    ],
    important='Key difference from BDY-E: The surveyor must establish corners using professional judgment. The AI cannot make this decision - it can only prepare the research.',
)

# ── FT ──
add_workflow(
    'Workflow 3: FT - FEMA Topographic Survey',
    'Elevation certificate or flood zone determination. No boundary work, no adjoiners.',
    [
        ['1', 'AI - Create project folders', 'AI', ''],
        ['2', 'AI - Search client deed', 'AI', 'Need property description for cert'],
        ['3', 'REVIEW - Verify correct parcel', 'Human', ''],
        ['4', 'AI - Pull flood zone data', 'AI', 'Query FEMA NFHL layers'],
        ['5', 'SKIP - No adjoiners needed', '-', 'Not needed for topo'],
    ],
    [
        ['6', 'FIELD - Set up GPS/total station', 'Surveyor', ''],
        ['7', 'FIELD - Establish benchmark', 'Surveyor', 'NAVD88 datum'],
        ['8', 'FIELD - Shoot topo points', 'Surveyor', 'Spot elevations, contours'],
        ['9', 'FIELD - Locate structures', 'Surveyor', 'LAG, HAG, next higher floor'],
        ['10', 'FIELD - Photograph reference marks', 'Surveyor', 'Required for FEMA form'],
    ],
    [
        ['11', 'AI - Generate surface / contours', 'AI (Civil 3D)', 'TIN surface, 1 ft contours'],
        ['12', 'HUMAN - FEMA Elevation Certificate', 'PLS', 'Professional form'],
        ['13', 'HUMAN - Sign & seal', 'PLS', ''],
    ],
)

# ── SUB ──
add_workflow(
    'Workflow 4: SUB - Subdivision',
    'Divide one parcel into multiple lots. Requires perimeter adjoiners, lot layout, and recordable plat.',
    [
        ['1', 'AI - Create project folders', 'AI', ''],
        ['2', 'AI - Search parent parcel deed', 'AI', 'The deed being subdivided'],
        ['3', 'REVIEW - Check restrictions/easements', 'Human', 'Zoning, covenants, setbacks'],
        ['4', 'AI - Search parent plat', 'AI', 'Existing plat of parent'],
        ['5', 'AI - Discover perimeter adjoiners', 'AI', 'Only surrounding parcels'],
        ['6', 'AI - Bulk research adjoiners', 'AI', 'Deeds + plats'],
        ['7', 'HUMAN - Check county regulations', 'PLS', 'Lot size mins, road requirements'],
        ['8', 'HUMAN - Design lot layout', 'PLS/Eng', 'Road widths, dimensions, utilities'],
    ],
    [
        ['9', 'FIELD - Survey parent boundary', 'Surveyor', 'Full perimeter'],
        ['10', 'FIELD - Topo of subdivision area', 'Surveyor', 'Drainage, utilities, roads'],
        ['11', 'FIELD - Set all new lot corners', 'Surveyor', 'Per approved layout'],
        ['12', 'FIELD - Set road/ROW monuments', 'Surveyor', ''],
    ],
    [
        ['13', 'AI+HUMAN - Subdivision plat', 'AI+Human', 'Lot layout, dimensions, areas'],
        ['14', 'AI - Insert reference table', 'AI', ''],
        ['15', 'HUMAN - Utility easements/setbacks', 'PLS/Eng', ''],
        ['16', 'HUMAN - County planning review', 'PLS', 'Submit for approval'],
        ['17', 'HUMAN - Certify, sign & seal', 'PLS', ''],
        ['18', 'HUMAN - Record with county clerk', 'PLS', ''],
    ],
)

# ── SE ──
add_workflow(
    'Workflow 5: SE - Survey Exam (Desktop Review)',
    'Review existing survey work for accuracy, completeness, or dispute resolution. No new field work.',
    [
        ['1', 'AI - Create project folders', 'AI', ''],
        ['2', 'AI - Pull ALL deeds in chain of title', 'AI', 'Deep chain search'],
        ['3', 'AI - Pull ALL plats (client + adjoiners)', 'AI', 'Full cabinet search'],
        ['4', 'AI - Full adjoiner research', 'AI', 'Every neighboring parcel'],
        ['5', 'AI - Analyze descriptions for conflicts', 'AI', 'Flag overlaps, gaps'],
        ['6', 'AI - Similarity search on legal desc', 'AI', 'Find related parcels via AI'],
        ['7', 'HUMAN - Professional analysis', 'PLS', 'Written opinion'],
        ['8', 'HUMAN - Generate written report', 'PLS', 'Findings and recommendations'],
    ],
    [],
    [],
    tip='Survey Exams are the ideal AI showcase - the entire job is research. The AI can do 80%+ of the work, with the surveyor providing professional analysis.',
)

# ── Discussion Section ──
doc.add_heading('Team Discussion Questions', level=1)
questions = [
    'Research depth per job type - Are the levels right? Should topo surveys pull any deed research at all?',
    'Auto-select confidence - Currently score >= 25 to auto-pick a deed. Too aggressive? Too conservative? Should adjoiners have a lower threshold than the client?',
    'Field protocols - What is the standard monument search order? Corner numbering convention? Photo naming?',
    'QA gates - Should someone review AI-selected deeds before field day? Or is post-field review sufficient?',
    'Deliverables - Standard deliverable sets per job type? (e.g., BDY = plat + legal desc, FT = elevation cert + topo DWG)',
    'County specifics - Different counties have different portals and filing requirements. Should the workflow adapt by county?',
    'Pricing integration - Should the inquiry form show an automatic estimate based on ML complexity prediction?',
    'Client communication - Should the system auto-email the client with a tracking link when research starts?',
]
for i, q in enumerate(questions, 1):
    doc.add_paragraph(f'{i}. {q}', style='List Number')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('--- Red Tail Surveying - Taos County, NM ---')
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.font.size = Pt(10)

# Save
out_path = r'J:\AI DATA CENTER\AI Agents\Deed & Plat Helper\Red_Tail_Workflow_Reference.docx'
doc.save(out_path)
print(f'Saved to: {out_path}')
